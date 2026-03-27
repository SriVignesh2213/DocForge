import logging
import os
from docx import Document
from docxcompose.composer import Composer
from doc_utils import iter_block_items, replace_body_section_properties
from docx.text.paragraph import Paragraph
from docx.table import Table
import re
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from table_handler import TableHandler
from image_handler import ImageHandler

logger = logging.getLogger(__name__)

class Formatter:
    HEADING_BLUE = RGBColor(0x00, 0x70, 0xC0)
    NUMBERING_RE = re.compile(r'^\s*(\d+(?:\.\d+)*)(?:[\.\)])?\s+')
    TABLE_CAPTION_RE = re.compile(r'^\s*(?:\d+(?:\.\d+)*\.?\s+)?table\b', re.IGNORECASE)
    FIGURE_CAPTION_RE = re.compile(r'^\s*(?:\d+(?:\.\d+)*\.?\s+)?(?:figure|fig)\.?\b', re.IGNORECASE)
    MAJOR_SECTION_TERMS = (
        "abstract",
        "introduction",
        "background",
        "literature review",
        "materials and methods",
        "materials & methods",
        "method",
        "methodology",
        "results",
        "results and discussion",
        "discussion",
        "conclusion",
        "conclusions",
        "acknowledgement",
        "acknowledgements",
        "references",
        "authors' note",
        "authors’ note",
    )

    def __init__(self, template_path, input_path, output_path):
        self.template_path = template_path
        self.input_path = input_path
        self.output_path = output_path
        self.table_handler = TableHandler()
        self.image_handler = ImageHandler()

    def _coerce_int(self, value, default=0):
        if value is None:
            return default

        try:
            return int(value)
        except (TypeError, ValueError):
            try:
                return int(round(float(str(value))))
            except (TypeError, ValueError):
                logger.warning(f"Falling back to default integer for unsupported value: {value!r}")
                return default

    def _get_paragraph_metrics(self, paragraph):
        max_size = 0
        is_bold = False

        for run in paragraph.runs:
            if run.bold:
                is_bold = True
            if run.font.size and run.font.size.pt > max_size:
                max_size = run.font.size.pt

        return is_bold, max_size

    def _is_italic_paragraph(self, paragraph):
        return any(run.italic for run in paragraph.runs if run.text.strip())

    def _extract_numbering(self, text):
        match = self.NUMBERING_RE.match(text)
        return match.group(0) if match else ""

    def _strip_numbering(self, text):
        return self.NUMBERING_RE.sub("", text, count=1).strip()

    def _normalized_heading_text(self, text):
        return self._strip_numbering(text).strip().lower()

    def _is_non_section_heading(self, text):
        normalized = self._normalized_heading_text(text)
        return (
            normalized.startswith("keywords:")
            or normalized.startswith("figure labels:")
            or normalized.startswith("figure label:")
            or normalized.startswith("table labels:")
            or normalized.startswith("table label:")
            or self.TABLE_CAPTION_RE.match(text)
            or self.FIGURE_CAPTION_RE.match(text)
            or normalized.startswith("chart ")
            or normalized.startswith("chart.")
            or normalized.startswith("scheme ")
            or normalized.startswith("scheme.")
        )

    def _get_list_level(self, paragraph):
        paragraph_properties = paragraph._p.pPr
        if paragraph_properties is None:
            return None

        numbering_properties = paragraph_properties.find(qn('w:numPr'))
        if numbering_properties is None:
            return None

        level = numbering_properties.find(qn('w:ilvl'))
        if level is None:
            return 0

        return self._coerce_int(level.get(qn('w:val')), 0)

    def _is_heading_candidate(self, paragraph, text, is_bold, max_size):
        style_name = paragraph.style.name.lower()
        list_level = self._get_list_level(paragraph)
        return (
            ('heading' in style_name)
            or (is_bold and 10 < max_size <= 14 and len(text.split()) < 20)
            or (self._extract_numbering(text) and len(text.split()) < 20)
            or (list_level is not None and list_level >= 1 and len(text.split()) < 20)
        )

    def _is_section_heading(self, paragraph, text, is_bold, max_size):
        if self._is_non_section_heading(text):
            return False

        normalized = self._normalized_heading_text(text)
        if normalized in self.MAJOR_SECTION_TERMS:
            return True
        if self._is_heading_candidate(paragraph, text, is_bold, max_size):
            return True

        letters = [char for char in self._strip_numbering(text) if char.isalpha()]
        return bool(
            letters
            and len(text.split()) < 20
            and sum(char.isupper() for char in letters) >= max(3, int(len(letters) * 0.6))
        )

    def _is_front_matter_boundary(self, paragraph, text, is_bold, max_size):
        if self._is_non_section_heading(text):
            return False

        normalized = self._normalized_heading_text(text)
        if normalized in self.MAJOR_SECTION_TERMS:
            return True

        if self._is_italic_paragraph(paragraph):
            return False

        if self._extract_numbering(text) and len(text.split()) < 20:
            return True

        letters = [char for char in self._strip_numbering(text) if char.isalpha()]
        if letters and len(text.split()) < 20 and sum(char.isupper() for char in letters) >= max(3, int(len(letters) * 0.6)):
            return True

        return is_bold and 10 < max_size <= 14 and len(text.split()) < 12

    def _get_heading_level(self, paragraph, text):
        style_name = paragraph.style.name.lower()
        digits = ''.join(filter(str.isdigit, style_name))
        if digits:
            return 2 if int(digits) >= 2 else 1

        numbering = self._extract_numbering(text).strip()
        if numbering:
            numbering = numbering.rstrip('.)')
            return 2 if '.' in numbering else 1

        list_level = self._get_list_level(paragraph)
        if list_level is not None:
            return 1 if list_level == 0 else 2

        normalized = self._normalized_heading_text(text)
        if any(term in normalized for term in self.MAJOR_SECTION_TERMS):
            return 1

        if "sub heading" in normalized or "subheading" in normalized or "sub section" in normalized or "subsection" in normalized:
            return 2

        letters = [char for char in self._strip_numbering(text) if char.isalpha()]
        if letters and sum(char.isupper() for char in letters) >= max(3, int(len(letters) * 0.6)):
            return 1

        return 2

    def _replace_paragraph_text(self, paragraph, new_text):
        if not paragraph.runs:
            paragraph.add_run(new_text)
            return

        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""

    def _has_list_numbering(self, paragraph):
        paragraph_properties = paragraph._p.pPr
        return paragraph_properties is not None and paragraph_properties.find(qn('w:numPr')) is not None

    def _clear_list_numbering(self, paragraph):
        paragraph_properties = paragraph._p.pPr
        if paragraph_properties is None:
            return

        numbering_properties = paragraph_properties.find(qn('w:numPr'))
        if numbering_properties is not None:
            paragraph_properties.remove(numbering_properties)

    def _sync_numbering_counters(self, numbering_text, level, h1_num, h2_num):
        parts = [self._coerce_int(part, 0) for part in numbering_text.rstrip('.)').split('.')]
        if not parts:
            return h1_num, h2_num

        if level == 1 or len(parts) == 1:
            return parts[0], 0

        return parts[0], parts[1]

    def _get_section_column_count(self, sectpr):
        columns = sectpr.find(qn('w:cols'))
        if columns is None:
            return 1
        return max(1, self._coerce_int(columns.get(qn('w:num')), 1))

    def _build_layout_profile(self, template_doc):
        body_section = template_doc.sections[-1]
        for section in reversed(template_doc.sections):
            if self._get_section_column_count(section._sectPr) == 2:
                body_section = section
                break

        available_width = (
            self._coerce_int(body_section.page_width)
            - self._coerce_int(body_section.left_margin)
            - self._coerce_int(body_section.right_margin)
        )
        column_count = max(1, self._get_section_column_count(body_section._sectPr))
        columns = body_section._sectPr.find(qn('w:cols'))
        spacing_twips = self._coerce_int(columns.get(qn('w:space')), 0) if columns is not None else 0
        spacing = spacing_twips * 635
        column_width = self._coerce_int((available_width - (spacing * (column_count - 1))) / column_count)

        return {
            "double_section_sectpr": body_section._sectPr,
            "double_column_count": column_count,
            "double_column_width": column_width,
            "single_column_width": available_width,
        }

    def _extract_template_run_profile(self, paragraph, default_size, default_bold=False, default_italic=False, default_color=None):
        style_font = paragraph.style.font if paragraph.style is not None and hasattr(paragraph.style, "font") else None

        font_name = None
        font_size = None
        color = None
        bold = default_bold
        italic = default_italic

        for run in paragraph.runs:
            if not run.text.strip():
                continue
            if font_name is None and run.font.name:
                font_name = run.font.name
            if font_size is None and run.font.size:
                font_size = run.font.size.pt
            if color is None and run.font.color and run.font.color.rgb:
                color = run.font.color.rgb
            if run.bold:
                bold = True
            if run.italic:
                italic = True

        if style_font is not None:
            if font_name is None and style_font.name:
                font_name = style_font.name
            if font_size is None and style_font.size:
                font_size = style_font.size.pt
            if color is None and style_font.color and style_font.color.rgb:
                color = style_font.color.rgb
            if style_font.bold:
                bold = True
            if style_font.italic:
                italic = True

        return {
            "font_name": font_name or "Times New Roman",
            "font_size": font_size or default_size,
            "bold": bold,
            "italic": italic,
            "color": color if color is not None else default_color,
        }

    def _build_template_text_profile(self, template_doc):
        template_paragraphs = [paragraph for paragraph in template_doc.paragraphs if paragraph.text.strip()]
        if not template_paragraphs:
            return {
                "title": {"font_name": "Times New Roman", "font_size": 16, "bold": True, "italic": False, "color": self.HEADING_BLUE},
                "front_matter": {"font_name": "Times New Roman", "font_size": 12, "bold": False, "italic": True, "color": RGBColor(0, 0, 0)},
                "heading": {"font_name": "Times New Roman", "font_size": 12, "bold": True, "italic": False, "color": self.HEADING_BLUE},
                "body": {"font_name": "Times New Roman", "font_size": 12, "bold": False, "italic": False, "color": RGBColor(0, 0, 0)},
            }

        title_paragraph = template_paragraphs[0]
        heading_index = None
        heading_paragraph = None

        for index, paragraph in enumerate(template_paragraphs[1:], start=1):
            text = paragraph.text.strip()
            is_bold, max_size = self._get_paragraph_metrics(paragraph)
            if self._is_section_heading(paragraph, text, is_bold, max_size):
                heading_index = index
                heading_paragraph = paragraph
                break

        front_paragraph = None
        if heading_index is not None:
            for paragraph in template_paragraphs[1:heading_index]:
                if paragraph.text.strip():
                    front_paragraph = paragraph
                    break
        elif len(template_paragraphs) > 1:
            front_paragraph = template_paragraphs[1]

        body_paragraph = None
        search_start = heading_index + 1 if heading_index is not None else 1
        for paragraph in template_paragraphs[search_start:]:
            text = paragraph.text.strip()
            if not text or self._is_non_section_heading(text):
                continue
            is_bold, max_size = self._get_paragraph_metrics(paragraph)
            if self._is_section_heading(paragraph, text, is_bold, max_size):
                continue
            body_paragraph = paragraph
            break

        heading_profile = self._extract_template_run_profile(
            heading_paragraph or title_paragraph,
            12,
            default_bold=True,
            default_color=self.HEADING_BLUE,
        )
        title_profile = self._extract_template_run_profile(
            title_paragraph,
            16,
            default_bold=True,
            default_color=heading_profile["color"],
        )
        front_profile = self._extract_template_run_profile(
            front_paragraph or title_paragraph,
            12,
            default_italic=True,
            default_color=RGBColor(0, 0, 0),
        )
        body_profile = self._extract_template_run_profile(
            body_paragraph or front_paragraph or title_paragraph,
            12,
            default_color=RGBColor(0, 0, 0),
        )
        body_profile["bold"] = False
        body_profile["italic"] = False
        front_profile["bold"] = False
        front_profile["italic"] = True
        heading_profile["bold"] = True
        heading_profile["italic"] = False

        return {
            "title": title_profile,
            "front_matter": front_profile,
            "heading": heading_profile,
            "body": body_profile,
        }

    def _apply_text_profile(self, paragraph, profile, alignment=None, force_bold=None, force_italic=None, force_color=None):
        if alignment is not None:
            paragraph.alignment = alignment

        color = force_color if force_color is not None else profile.get("color")
        for run in paragraph.runs:
            run.font.name = profile.get("font_name") or "Times New Roman"
            if profile.get("font_size"):
                run.font.size = Pt(profile["font_size"])
            run.font.bold = profile["bold"] if force_bold is None else force_bold
            run.font.italic = profile["italic"] if force_italic is None else force_italic
            if color is not None:
                run.font.color.rgb = color

    def _apply_template_text_formatting(self, document, template_text_profile):
        h1_num = 0
        h2_num = 0
        seen_first_section_heading = False

        paragraphs = [block for block in iter_block_items(document) if isinstance(block, Paragraph)]
        title_paragraph = next((paragraph for paragraph in paragraphs if paragraph.text.strip()), None)
        first_section_heading_paragraph = None
        has_abstract_heading = False
        preserve_uploaded_numbering = False

        if title_paragraph is not None:
            title_found = False
            for paragraph in paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                if not title_found:
                    title_found = paragraph is title_paragraph
                    continue

                is_bold, max_size = self._get_paragraph_metrics(paragraph)
                if self._is_front_matter_boundary(paragraph, text, is_bold, max_size):
                    first_section_heading_paragraph = paragraph
                    break

        for paragraph in paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            is_bold, max_size = self._get_paragraph_metrics(paragraph)
            if not self._is_section_heading(paragraph, text, is_bold, max_size):
                continue

            normalized = self._normalized_heading_text(text)
            if normalized == "abstract":
                has_abstract_heading = True

            if self._extract_numbering(text):
                preserve_uploaded_numbering = True
                break

        numbering_started = not has_abstract_heading
        numbering_stopped = False

        for paragraph in paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            for run in paragraph.runs:
                if run.text:
                    run.text = re.sub(r'\[\d+\]', '', run.text)

            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            is_bold, max_size = self._get_paragraph_metrics(paragraph)
            is_heading_candidate = self._is_heading_candidate(paragraph, text, is_bold, max_size)
            is_section_heading = self._is_section_heading(paragraph, text, is_bold, max_size)

            if paragraph is title_paragraph:
                self._apply_text_profile(
                    paragraph,
                    template_text_profile["title"],
                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    force_bold=True,
                    force_italic=False,
                )
                continue

            if not seen_first_section_heading and paragraph is not first_section_heading_paragraph:
                self._apply_text_profile(
                    paragraph,
                    template_text_profile["front_matter"],
                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    force_bold=False,
                    force_italic=True,
                    force_color=RGBColor(0, 0, 0),
                )
                continue

            if is_heading_candidate and len(text.split()) < 20:
                seen_first_section_heading = True
                if is_section_heading:
                    normalized = self._normalized_heading_text(text)
                    level = self._get_heading_level(paragraph, text)
                    explicit_numbering = self._extract_numbering(text).strip()
                    has_list_numbering = self._has_list_numbering(paragraph)

                    if normalized == "abstract":
                        numbering_started = True
                    elif "conclusion" in normalized or "references" in normalized:
                        numbering_stopped = True

                    if explicit_numbering:
                        h1_num, h2_num = self._sync_numbering_counters(explicit_numbering, level, h1_num, h2_num)
                    elif numbering_started and not numbering_stopped and normalized != "abstract":
                        if level == 1:
                            h1_num += 1
                            h2_num = 0
                            if has_list_numbering or not preserve_uploaded_numbering:
                                self._replace_paragraph_text(paragraph, f"{h1_num}. {self._strip_numbering(text)}")
                        else:
                            if h1_num == 0:
                                h1_num = 1
                            h2_num += 1
                            if has_list_numbering or not preserve_uploaded_numbering:
                                self._replace_paragraph_text(paragraph, f"{h1_num}.{h2_num}. {self._strip_numbering(text)}")

                    if has_list_numbering:
                        self._clear_list_numbering(paragraph)

                self._apply_text_profile(
                    paragraph,
                    template_text_profile["heading"],
                    alignment=WD_ALIGN_PARAGRAPH.LEFT,
                    force_bold=True,
                    force_italic=False,
                )
                paragraph.paragraph_format.left_indent = Pt(0)
                paragraph.paragraph_format.first_line_indent = Pt(0)
                continue

            self._apply_text_profile(
                paragraph,
                template_text_profile["body"],
                alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                force_bold=False,
                force_italic=False,
                force_color=RGBColor(0, 0, 0),
            )

    def apply_styles_and_build(self, input_sections, mapping):
        input_doc = Document(self.input_path)
        
        # 1. Update styles in input document based on mapping
        for in_sec in input_sections:
            target_template = mapping.get(in_sec.title)
            if not target_template:
                continue
                
            for element in in_sec.elements:
                # Need to find the corresponding element in the loaded input_doc
                # For simplicity, docx compose keeps the objects in order
                pass
                
        # A more direct approach to modifying the input doc styles:
        current_mapped_template = None
        for block in iter_block_items(input_doc):
            if isinstance(block, Paragraph):
                style_name = block.style.name.lower()
                if 'heading' in style_name:
                    title = block.text.strip()
                    if title in mapping:
                        current_mapped_template = mapping[title]
                        try:
                            block.style = current_mapped_template.heading_style
                        except:
                            pass
                else:
                    if current_mapped_template:
                        try:
                            block.style = current_mapped_template.body_style
                        except:
                            pass
            elif isinstance(block, Table):
                self.table_handler.apply_template_table_style(block)
                
        template_doc = Document(self.template_path)
        template_text_profile = self._build_template_text_profile(template_doc)
        self._apply_template_text_formatting(input_doc, template_text_profile)

        # Rebuild using Composer to inherit template master properties
        layout_profile = self._build_layout_profile(template_doc)

        styled_input_path = "styled_temp.docx"
        input_doc.save(styled_input_path)
        
        # Merge input doc into the empty template envelope natively
        
        # 3. Flawlessly strip template dummy bytes without killing its multi-section layout logic (where the header logos live!)
        for p in template_doc.paragraphs:
            if p._element.xpath('.//w:sectPr'):
                p.clear() # DO NOT delete node, it anchors an official layout section break!
            else:
                p._element.getparent().remove(p._element) # Safe to permanently delete dummy placeholder
                
        for t in template_doc.tables:
            t._element.getparent().remove(t._element) # Delete dummy tables
            
        composer = Composer(template_doc)
        styled_doc = Document(styled_input_path)
        composer.append(styled_doc)
        
        composer.save(self.output_path)

        final_doc = Document(self.output_path)
        self._apply_template_text_formatting(final_doc, template_text_profile)
        wrapped_large_tables = self.table_handler.optimize_table_layout(final_doc, layout_profile)
        wrapped_large_figures = self.image_handler.optimize_figure_layout(final_doc, layout_profile)
        if wrapped_large_tables or wrapped_large_figures:
            replace_body_section_properties(
                final_doc,
                layout_profile["double_section_sectpr"],
                layout_profile["double_column_count"],
            )

        self.image_handler.validate_images(final_doc)
        final_doc.save(self.output_path)
        logger.info(f"Successfully generated formatted document {self.output_path}")
        
        if os.path.exists(styled_input_path):
            os.remove(styled_input_path)
